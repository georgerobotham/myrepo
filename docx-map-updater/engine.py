"""
docx-map-updater engine
------------------------
Pure logic (no GUI/tkinter dependency): finds a figures section in a Word
report, parses its existing figures, matches a folder of image files against
them, and applies replacements/insertions. Import this from a GUI or a
console script.
"""

from __future__ import annotations

import copy
import io
import json
import random
import re
import shutil
from dataclasses import dataclass, field
from difflib import SequenceMatcher
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from PIL import Image

IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".tif", ".tiff", ".bmp"}
EMU_PER_INCH = 914400
ID_LOCAL_NAMES = {"paraId", "anchorId", "editId"}
ID_PATTERN = re.compile(r"^[0-9A-F]{8}$")
CONFLICT_EPSILON = 0.05


# ---------------------------------------------------------------------------
# Data model
# ---------------------------------------------------------------------------

@dataclass
class Figure:
    number: int
    caption_text: str
    image_para: object
    caption_para: object
    embed_rid: str | None
    extent_cx: int | None
    extent_cy: int | None


@dataclass
class MatchPlan:
    section_heading: str
    replacements: list = field(default_factory=list)
    insertions: list = field(default_factory=list)
    unmatched_figures: list = field(default_factory=list)
    conflicts: list = field(default_factory=list)

    def to_report_dict(self) -> dict:
        return {
            "section_heading": self.section_heading,
            "replacements": [
                {
                    "figure_number": r["figure"].number,
                    "file": str(r["file"]),
                    "score": round(r["score"], 4),
                    "low_confidence": r["score"] < 0.75,
                }
                for r in self.replacements
            ],
            "insertions": [
                {
                    "figure_number": i["number"],
                    "file": str(i["file"]),
                    "suggested_caption": i["caption"],
                }
                for i in self.insertions
            ],
            "unmatched_figures": [
                {"figure_number": f.number, "caption": f.caption_text}
                for f in self.unmatched_figures
            ],
            "conflicts": [
                {
                    "figure_number": c["figure"].number,
                    "caption": c["figure"].caption_text,
                    "candidates": [
                        {"file": str(f), "score": round(s, 4)} for f, s in c["candidates"]
                    ],
                }
                for c in self.conflicts
            ],
        }


class ValidationError(RuntimeError):
    pass


# ---------------------------------------------------------------------------
# Step 1-2: locate section, parse existing figures
# ---------------------------------------------------------------------------

def _section_bounds_from_index(document, heading_idx: int):
    """Given the index of a heading paragraph, return (heading_idx,
    section_paragraph_indices), using that paragraph's own style as the
    section boundary marker: the section runs until the next paragraph
    using that same style, or end of document."""
    paragraphs = document.paragraphs
    heading_style = paragraphs[heading_idx].style.name
    end_idx = len(paragraphs)
    for i in range(heading_idx + 1, len(paragraphs)):
        if paragraphs[i].style.name == heading_style:
            end_idx = i
            break

    section_indices = list(range(heading_idx + 1, end_idx))
    return heading_idx, section_indices


def find_section_paragraphs(document, heading_text: str):
    """Return (heading_para_index, section_paragraph_indices), locating the
    heading by a case-insensitive substring match against paragraph text."""
    paragraphs = document.paragraphs
    heading_needle = heading_text.strip().lower()
    if not heading_needle:
        raise ValueError("--heading must not be empty")

    heading_idx = None
    for i, p in enumerate(paragraphs):
        if heading_needle in p.text.strip().lower():
            heading_idx = i
            break
    if heading_idx is None:
        raise ValueError(f"No paragraph found containing heading text: {heading_text!r}")

    return _section_bounds_from_index(document, heading_idx)


def find_section_by_index(document, heading_idx: int):
    """Return (heading_para_index, section_paragraph_indices) for an
    explicitly chosen heading paragraph index (e.g. from list_headings)."""
    paragraphs = document.paragraphs
    if heading_idx < 0 or heading_idx >= len(paragraphs):
        raise ValueError(f"heading_index {heading_idx} is out of range")
    return _section_bounds_from_index(document, heading_idx)


HEADING_STYLE_RE = re.compile(r"^(heading|title)", re.IGNORECASE)


def list_headings(document) -> list[dict]:
    """List paragraphs using a heading-like style, for a UI section picker."""
    headings = []
    for i, p in enumerate(document.paragraphs):
        text = p.text.strip()
        if not text:
            continue
        style_name = p.style.name if p.style is not None else ""
        if HEADING_STYLE_RE.match(style_name):
            headings.append({"index": i, "text": text, "style": style_name})
    return headings


CAPTION_RE_TEMPLATE = r"^{prefix}\s+(\d+):\s*(.*)$"


def _para_has_drawing(paragraph) -> bool:
    return paragraph._p.find(".//" + qn("w:drawing")) is not None


def _find_blip_rid(paragraph):
    blip = paragraph._p.find(".//" + qn("a:blip"))
    if blip is None:
        return None
    return blip.get(qn("r:embed"))


def _find_extent(paragraph):
    extent = paragraph._p.find(".//" + qn("wp:extent"))
    if extent is None:
        return None, None
    cx = extent.get("cx")
    cy = extent.get("cy")
    return (int(cx) if cx else None, int(cy) if cy else None)


def parse_existing_figures(document, section_indices, figure_prefix: str) -> list[Figure]:
    paragraphs = document.paragraphs
    caption_re = re.compile(CAPTION_RE_TEMPLATE.format(prefix=re.escape(figure_prefix)))
    figures: list[Figure] = []

    i = 0
    while i < len(section_indices):
        idx = section_indices[i]
        para = paragraphs[idx]
        if _para_has_drawing(para):
            j = i + 1
            while j < len(section_indices) and not paragraphs[section_indices[j]].text.strip():
                j += 1
            if j < len(section_indices):
                caption_para = paragraphs[section_indices[j]]
                m = caption_re.match(caption_para.text.strip())
                if m:
                    number = int(m.group(1))
                    caption_text = m.group(2).strip()
                    cx, cy = _find_extent(para)
                    figures.append(
                        Figure(
                            number=number,
                            caption_text=caption_text,
                            image_para=para,
                            caption_para=caption_para,
                            embed_rid=_find_blip_rid(para),
                            extent_cx=cx,
                            extent_cy=cy,
                        )
                    )
                    i = j + 1
                    continue
        i += 1

    return figures


# ---------------------------------------------------------------------------
# Step 3: match image files to existing figures
# ---------------------------------------------------------------------------

TOKEN_RE = re.compile(r"[^a-z0-9%]+")
# Case-preserving variant, used where original casing matters (e.g. keeping
# an acronym like "STC" recognisable when suggesting a caption).
TOKEN_SPLIT_RE = re.compile(r"[^A-Za-z0-9%]+")


def tokenize(text: str) -> set[str]:
    text = text.lower()
    tokens = [t for t in TOKEN_RE.split(text) if t]
    return set(tokens)


def score_pair(file_tokens: set[str], caption_tokens: set[str]) -> float:
    if not file_tokens and not caption_tokens:
        return 0.0
    union = file_tokens | caption_tokens
    jaccard = len(file_tokens & caption_tokens) / len(union) if union else 0.0
    seq_ratio = SequenceMatcher(
        None, " ".join(sorted(file_tokens)), " ".join(sorted(caption_tokens))
    ).ratio()
    return 0.5 * jaccard + 0.5 * seq_ratio


def list_image_files(images_dir: Path) -> list[Path]:
    return sorted(
        p for p in images_dir.iterdir()
        if p.is_file() and p.suffix.lower() in IMAGE_EXTENSIONS
    )


# Trailing filename tokens that describe file/version status rather than
# content, stripped when suggesting a caption (case-insensitive).
CAPTION_STRIP_SUFFIXES = {
    "final", "updated", "update", "new", "draft", "revised", "rev",
    "edited", "edit", "copy", "latest",
}
CAPTION_VERSION_RE = re.compile(r"^v\d+$", re.IGNORECASE)


def suggest_caption(filename_stem: str, general_style: str = "") -> str:
    """Heuristic, fully-offline caption suggestion for a new figure: clean
    up the filename (strip version/status tokens, underscores -> spaces,
    title case) and optionally fold in a user-supplied general caption
    style/note. Always presented to the user for review before it's used."""
    tokens = [t for t in TOKEN_SPLIT_RE.split(filename_stem) if t]
    trimmed = list(tokens)
    while trimmed and (trimmed[-1].lower() in CAPTION_STRIP_SUFFIXES or CAPTION_VERSION_RE.match(trimmed[-1])):
        trimmed.pop()
    if trimmed:
        tokens = trimmed

    words = []
    for t in tokens:
        # Preserve already-uppercase acronyms (e.g. "STC"); title-case the rest.
        words.append(t if t.isupper() and len(t) > 1 else t.capitalize())
    cleaned = " ".join(words) or filename_stem

    general_style = general_style.strip()
    if general_style:
        return f"{cleaned} - {general_style}"
    return cleaned


def build_match_plan(
    section_heading: str,
    figures: list[Figure],
    image_files: list[Path],
    threshold: float,
    general_style: str = "",
) -> MatchPlan:
    plan = MatchPlan(section_heading=section_heading)

    file_tokens = {f: tokenize(f.stem) for f in image_files}
    fig_tokens = {fig.number: tokenize(fig.caption_text) for fig in figures}
    figures_by_number = {fig.number: fig for fig in figures}

    pairs = []
    for f in image_files:
        for fig in figures:
            s = score_pair(file_tokens[f], fig_tokens[fig.number])
            pairs.append((s, f, fig.number))
    pairs.sort(key=lambda t: t[0], reverse=True)

    claimed_files: set[Path] = set()
    claimed_figures: set[int] = set()
    conflicted_figures: set[int] = set()

    # scores per figure, sorted desc, for conflict detection
    scores_by_figure: dict[int, list[tuple[float, Path]]] = {}
    for s, f, fig_num in pairs:
        scores_by_figure.setdefault(fig_num, []).append((s, f))

    for s, f, fig_num in pairs:
        if s < threshold:
            continue
        if f in claimed_files or fig_num in claimed_figures or fig_num in conflicted_figures:
            continue

        candidates = [
            (cs, cf) for cs, cf in scores_by_figure[fig_num]
            if cf not in claimed_files and cs >= threshold
        ]
        top_score = candidates[0][0]
        near_top = [c for c in candidates if top_score - c[0] <= CONFLICT_EPSILON]
        if len(near_top) > 1:
            conflicted_figures.add(fig_num)
            plan.conflicts.append(
                {
                    "figure": figures_by_number[fig_num],
                    "candidates": [(cf, cs) for cs, cf in near_top],
                }
            )
            continue

        claimed_files.add(f)
        claimed_figures.add(fig_num)
        plan.replacements.append({"figure": figures_by_number[fig_num], "file": f, "score": s})

    max_number = max((fig.number for fig in figures), default=0)
    next_number = max_number + 1
    for f in image_files:
        if f in claimed_files:
            continue
        plan.insertions.append({
            "number": next_number, "file": f,
            "caption": suggest_caption(f.stem, general_style),
        })
        next_number += 1

    for fig in figures:
        if fig.number not in claimed_figures and fig.number not in conflicted_figures:
            plan.unmatched_figures.append(fig)

    return plan


# ---------------------------------------------------------------------------
# Image geometry helpers
# ---------------------------------------------------------------------------

def _image_aspect_ratio(image_bytes: bytes) -> float:
    with Image.open(io.BytesIO(image_bytes)) as img:
        px_w, px_h = img.size
        dpi_x, dpi_y = img.info.get("dpi", (96, 96))
        dpi_x = dpi_x or 96
        dpi_y = dpi_y or 96
    return (px_w / dpi_x) / (px_h / dpi_y)


# ---------------------------------------------------------------------------
# Step 5: apply replacements
# ---------------------------------------------------------------------------

def apply_replacement(document, figure: Figure, image_path: Path, report_entry: dict) -> None:
    new_bytes = image_path.read_bytes()
    rid = figure.embed_rid
    if rid is None:
        raise ValidationError(f"Figure {figure.number} has no image relationship to replace")

    image_part = document.part.related_parts[rid]
    image_part._blob = new_bytes

    if figure.extent_cx and figure.extent_cy:
        current_aspect = figure.extent_cx / figure.extent_cy
        new_aspect = _image_aspect_ratio(new_bytes)
        if abs(new_aspect - current_aspect) / current_aspect > 0.02:
            new_cy = round(figure.extent_cx / new_aspect)
            _set_extent(figure.image_para, figure.extent_cx, new_cy)
            report_entry["resized"] = True
            report_entry["new_extent_cy"] = new_cy


def _set_extent(paragraph, cx: int, cy: int) -> None:
    for tag in ("wp:extent",):
        el = paragraph._p.find(".//" + qn(tag))
        if el is not None:
            el.set("cx", str(cx))
            el.set("cy", str(cy))
    xfrm_ext = paragraph._p.find(".//" + qn("a:xfrm") + "/" + qn("a:ext"))
    if xfrm_ext is not None:
        xfrm_ext.set("cx", str(cx))
        xfrm_ext.set("cy", str(cy))


# ---------------------------------------------------------------------------
# Step 6: apply insertions
# ---------------------------------------------------------------------------

def _regenerate_ids(element) -> None:
    for el in element.iter():
        for attr in list(el.attrib):
            local_name = attr.split("}")[-1]
            if local_name in ID_LOCAL_NAMES:
                el.set(attr, "".join(random.choice("0123456789ABCDEF") for _ in range(8)))


def _set_caption_text(caption_para, text: str) -> None:
    runs = caption_para.runs
    if not runs:
        caption_para.add_run(text)
        return
    runs[0].text = text
    for extra in runs[1:]:
        extra.text = ""


def apply_insertion(
    document,
    template_figure: Figure,
    median_width_cx: int,
    number: int,
    image_path: Path,
    caption: str,
    figure_prefix: str,
    insert_before_para,
) -> None:
    new_bytes = image_path.read_bytes()
    rid, _image = document.part.get_or_add_image(io.BytesIO(new_bytes))

    image_para_copy = copy.deepcopy(template_figure.image_para._p)
    caption_para_copy = copy.deepcopy(template_figure.caption_para._p)
    _regenerate_ids(image_para_copy)
    _regenerate_ids(caption_para_copy)

    blip = image_para_copy.find(".//" + qn("a:blip"))
    if blip is not None:
        blip.set(qn("r:embed"), rid)

    new_aspect = _image_aspect_ratio(new_bytes)
    new_cy = round(median_width_cx / new_aspect)

    from docx.text.paragraph import Paragraph
    image_para = Paragraph(image_para_copy, template_figure.image_para._parent)
    caption_para = Paragraph(caption_para_copy, template_figure.caption_para._parent)
    _set_extent(image_para, median_width_cx, new_cy)
    _set_caption_text(caption_para, f"{figure_prefix} {number}: {caption}")

    anchor = insert_before_para._p
    anchor.addprevious(image_para_copy)
    anchor.addprevious(caption_para_copy)


# ---------------------------------------------------------------------------
# Step 7: save and validate
# ---------------------------------------------------------------------------

def validate_saved_docx(path: Path) -> None:
    doc = Document(str(path))
    bad = []
    for part in [doc.part] + list(doc.part.package.iter_parts()):
        if not hasattr(part, "element") or part.element is None:
            continue
        try:
            for el in part.element.iter():
                for attr, value in el.attrib.items():
                    local_name = attr.split("}")[-1]
                    if local_name in ID_LOCAL_NAMES and not ID_PATTERN.match(value):
                        bad.append((part.partname, el.tag, local_name, value))
        except Exception:
            continue
    if bad:
        details = "; ".join(f"{p} {tag} {name}={val!r}" for p, tag, name, val in bad[:10])
        raise ValidationError(f"Invalid ID attribute(s) found after save: {details}")


# ---------------------------------------------------------------------------
# Orchestration
# ---------------------------------------------------------------------------

def run(
    docx_path: Path,
    images_dir: Path,
    heading: str | None = None,
    heading_index: int | None = None,
    figure_prefix: str = "Figure",
    threshold: float = 0.55,
    apply: bool = False,
    in_place: bool = False,
    out_path: Path | None = None,
    report_path: Path | None = None,
    general_style: str = "",
    caption_overrides: dict[int, str] | None = None,
    log=print,
) -> dict:
    docx_path = Path(docx_path)
    images_dir = Path(images_dir)

    document = Document(str(docx_path))
    if heading_index is not None:
        heading_idx, section_indices = find_section_by_index(document, heading_index)
        heading = document.paragraphs[heading_idx].text.strip()
    elif heading:
        heading_idx, section_indices = find_section_paragraphs(document, heading)
    else:
        raise ValueError("Either heading or heading_index must be provided")

    figures = parse_existing_figures(document, section_indices, figure_prefix)
    if not figures:
        log(f"Warning: no existing figures found in section {heading!r}")

    image_files = list_image_files(images_dir)
    plan = build_match_plan(heading, figures, image_files, threshold, general_style=general_style)

    if caption_overrides:
        for ins in plan.insertions:
            override = caption_overrides.get(ins["number"])
            if override is not None and override.strip():
                ins["caption"] = override.strip()

    report = plan.to_report_dict()
    if report_path is None:
        report_path = docx_path.with_name(docx_path.stem + "_report.json")
    report_path.write_text(json.dumps(report, indent=2), encoding="utf-8")

    log(f"Section: {heading!r} - {len(figures)} existing figures, {len(image_files)} source images")
    log(f"  Replacements: {len(plan.replacements)}")
    log(f"  Insertions:   {len(plan.insertions)}")
    log(f"  Unmatched:    {len(plan.unmatched_figures)}")
    log(f"  Conflicts:    {len(plan.conflicts)}")
    log(f"Report written to {report_path}")

    if not apply:
        log("Dry run only - no changes written. Re-run with Apply to write changes.")
        return report

    for r in plan.replacements:
        entry = next(e for e in report["replacements"] if e["figure_number"] == r["figure"].number)
        apply_replacement(document, r["figure"], r["file"], entry)

    if plan.insertions:
        paragraphs = document.paragraphs
        last_figure = figures[-1] if figures else None
        if last_figure is None:
            raise ValidationError("Cannot insert new figures: section has no existing figures to use as a template")

        para_elements = [p._p for p in paragraphs]
        last_caption_idx = para_elements.index(last_figure.caption_para._p)
        insert_before = paragraphs[last_caption_idx + 1] if last_caption_idx + 1 < len(paragraphs) else None
        if insert_before is None:
            raise ValidationError("Cannot find insertion point after the last figure")

        widths = [f.extent_cx for f in figures if f.extent_cx]
        median_width = sorted(widths)[len(widths) // 2] if widths else 5486400  # 6in fallback

        for ins in plan.insertions:
            apply_insertion(
                document,
                last_figure,
                median_width,
                ins["number"],
                ins["file"],
                ins["caption"],
                figure_prefix,
                insert_before,
            )

    if out_path is None:
        if in_place:
            out_path = docx_path
        else:
            out_path = docx_path.with_name(docx_path.stem + "_updated.docx")

    if in_place:
        backup_path = docx_path.with_suffix(docx_path.suffix + ".bak")
        shutil.copy2(docx_path, backup_path)
        log(f"Backed up original to {backup_path}")

    document.save(str(out_path))
    log(f"Saved to {out_path}")

    validate_saved_docx(out_path)
    log("Validation passed: all paraId/anchorId/editId attributes are well-formed.")

    report["output_path"] = str(out_path)
    return report
