"""
Headless end-to-end test for engine.py: builds a synthetic .docx with known
figures plus real embedded images, runs the engine, and independently
verifies the output.
"""

import io
import json
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.shared import Emu, Inches
from PIL import Image

import engine

SCRATCH = Path(__file__).parent / "_test_scratch"


def make_image_bytes(px_w: int, px_h: int, color) -> bytes:
    img = Image.new("RGB", (px_w, px_h), color)
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


def build_fixture_docx(path: Path):
    doc = Document()
    doc.add_heading("Appendix A: Overview", level=1)
    doc.add_paragraph("Some intro text before the maps section.")

    doc.add_heading("Appendix B: STC and Erosion Maps", level=1)

    # Figure 1: 200x100 (aspect 2.0), width fixed at 2in
    f1_bytes = make_image_bytes(200, 100, (200, 0, 0))
    doc.add_picture(io.BytesIO(f1_bytes), width=Inches(2))
    doc.add_paragraph("Figure 1: STC Map North Zone")

    doc.add_paragraph("")  # blank paragraph to test skip-blank logic

    # Figure 2: 150x150 (aspect 1.0), width fixed at 2in
    f2_bytes = make_image_bytes(150, 150, (0, 200, 0))
    doc.add_picture(io.BytesIO(f2_bytes), width=Inches(2))
    doc.add_paragraph("Figure 2: Erosion Map South Zone")

    # Figure 3: 180x90 (aspect 2.0), width fixed at 2in
    f3_bytes = make_image_bytes(180, 90, (0, 0, 200))
    doc.add_picture(io.BytesIO(f3_bytes), width=Inches(2))
    doc.add_paragraph("Figure 3: STC Map East Zone")

    doc.add_heading("Appendix C: Conclusions", level=1)
    doc.add_paragraph("Some content after the maps section.")

    doc.save(str(path))
    return {"f1": f1_bytes, "f2": f2_bytes, "f3": f3_bytes}


def build_images_dir(path: Path):
    path.mkdir(parents=True, exist_ok=True)

    # Matches Figure 1 by tokens, same aspect (2.0) -> should keep frame size
    replace1 = make_image_bytes(400, 200, (250, 10, 10))
    (path / "stc_map_north_zone_v2.png").write_bytes(replace1)

    # Matches Figure 2 by tokens, different aspect (0.5 vs 1.0) -> should resize
    replace2 = make_image_bytes(150, 300, (10, 250, 10))
    (path / "erosion_map_south_zone_updated.png").write_bytes(replace2)

    # Matches nothing -> new insertion
    new_file = make_image_bytes(300, 100, (10, 10, 250))
    (path / "brand_new_culvert_diagram.png").write_bytes(new_file)

    return {
        "replace1": path / "stc_map_north_zone_v2.png",
        "replace2": path / "erosion_map_south_zone_updated.png",
        "new_file": path / "brand_new_culvert_diagram.png",
    }


def main():
    if SCRATCH.exists():
        shutil.rmtree(SCRATCH)
    SCRATCH.mkdir(parents=True)

    docx_path = SCRATCH / "sample_report.docx"
    images_dir = SCRATCH / "images"

    originals = build_fixture_docx(docx_path)
    files = build_images_dir(images_dir)

    failures = []

    def check(label, cond):
        status = "PASS" if cond else "FAIL"
        print(f"[{status}] {label}")
        if not cond:
            failures.append(label)

    # --- Dry run ---
    report_path = SCRATCH / "report.json"
    dry_report = engine.run(
        docx_path=docx_path,
        images_dir=images_dir,
        heading="STC and Erosion Maps",
        threshold=0.55,
        apply=False,
        report_path=report_path,
    )

    check("dry run: 2 replacements found", len(dry_report["replacements"]) == 2)
    check("dry run: 1 insertion found", len(dry_report["insertions"]) == 1)
    check("dry run: 1 unmatched figure (Figure 3)", len(dry_report["unmatched_figures"]) == 1)
    check("dry run: no output docx written",
          not (SCRATCH / "sample_report_updated.docx").exists())
    check("dry run: report.json written to disk", report_path.exists())

    replaced_numbers = {r["figure_number"] for r in dry_report["replacements"]}
    check("dry run: figures 1 and 2 are the replacements", replaced_numbers == {1, 2})

    unmatched_numbers = {f["figure_number"] for f in dry_report["unmatched_figures"]}
    check("dry run: figure 3 is unmatched", unmatched_numbers == {3})

    insertion = dry_report["insertions"][0]
    check("dry run: insertion numbered 4 (continues section sequence)", insertion["figure_number"] == 4)

    # --- Apply run ---
    apply_report = engine.run(
        docx_path=docx_path,
        images_dir=images_dir,
        heading="STC and Erosion Maps",
        threshold=0.55,
        apply=True,
        report_path=report_path,
    )
    out_path = Path(apply_report["output_path"])
    check("apply run: output file created", out_path.exists())
    check("apply run: original file untouched (non-destructive default)",
          docx_path.exists() and docx_path.stat().st_size > 0)

    # --- Independent verification of output ---
    out_doc = Document(str(out_path))
    paragraphs = out_doc.paragraphs
    texts = [p.text for p in paragraphs]

    check("output: pre-section content untouched",
          texts[0] == "Appendix A: Overview" and texts[1] == "Some intro text before the maps section.")
    check("output: post-section content untouched",
          "Appendix C: Conclusions" in texts and "Some content after the maps section." in texts)

    caption_texts = [t for t in texts if t.startswith("Figure ")]
    check("output: 4 figure captions present", len(caption_texts) == 4)
    check("output: figure 4 caption uses cleaned-up filename as suggested caption",
          any(t.startswith("Figure 4:") and "Brand New Culvert Diagram" in t for t in caption_texts))

    # Re-parse figures from the output and pixel-sample each image to confirm
    # the right bytes landed in the right slot.
    heading_idx, section_indices = engine.find_section_paragraphs(out_doc, "STC and Erosion Maps")
    out_figures = engine.parse_existing_figures(out_doc, section_indices, "Figure")
    check("output: 4 figures parsed back out", len(out_figures) == 4)

    figs_by_number = {f.number: f for f in out_figures}

    def image_bytes_for(fig):
        return out_doc.part.related_parts[fig.embed_rid].blob

    def pixel_matches(actual_bytes, expected_bytes):
        a = Image.open(io.BytesIO(actual_bytes)).convert("RGB")
        e = Image.open(io.BytesIO(expected_bytes)).convert("RGB")
        if a.size != e.size:
            return False
        return a.getpixel((a.width // 2, a.height // 2)) == e.getpixel((e.width // 2, e.height // 2))

    check("figure 1 now contains the replacement image (pixel sample)",
          pixel_matches(image_bytes_for(figs_by_number[1]), files["replace1"].read_bytes()))
    check("figure 2 now contains the replacement image (pixel sample)",
          pixel_matches(image_bytes_for(figs_by_number[2]), files["replace2"].read_bytes()))
    check("figure 3 (unmatched) still contains its original image (pixel sample)",
          pixel_matches(image_bytes_for(figs_by_number[3]), originals["f3"]))
    check("figure 4 (inserted) contains the new file's bytes (pixel sample)",
          pixel_matches(image_bytes_for(figs_by_number[4]), files["new_file"].read_bytes()))

    # Frame-sizing checks
    fig1 = figs_by_number[1]
    check("figure 1 frame size unchanged (replacement kept same aspect ratio)",
          fig1.extent_cx == Emu(Inches(2)) and fig1.extent_cy == Emu(Inches(1)))

    fig2 = figs_by_number[2]
    check("figure 2 frame width kept, height rescaled (aspect ratio changed)",
          fig2.extent_cx == Emu(Inches(2)) and fig2.extent_cy != Emu(Inches(2)))

    # ID validation (this is also called inside engine.run, but re-check explicitly)
    try:
        engine.validate_saved_docx(out_path)
        check("ID validation (paraId/anchorId/editId all well-formed)", True)
    except engine.ValidationError as exc:
        check(f"ID validation (paraId/anchorId/editId all well-formed): {exc}", False)

    # --- caption suggestion heuristic ---
    check("suggest_caption strips version/status suffixes",
          engine.suggest_caption("stc_map_north_zone_v2") == "Stc Map North Zone")
    check("suggest_caption preserves an already-uppercase acronym",
          engine.suggest_caption("STC_Map_North_updated") == "STC Map North")
    check("suggest_caption folds in a general style note",
          engine.suggest_caption("erosion_map_south", "2026 survey") == "Erosion Map South - 2026 survey")

    # --- heading picker (list_headings / find_section_by_index) ---
    scan_doc = Document(str(docx_path))
    headings = engine.list_headings(scan_doc)
    heading_texts = [h["text"] for h in headings]
    check("list_headings finds all 3 Heading-1 paragraphs",
          heading_texts == ["Appendix A: Overview", "Appendix B: STC and Erosion Maps", "Appendix C: Conclusions"])

    target = next(h for h in headings if h["text"] == "Appendix B: STC and Erosion Maps")
    idx_report = engine.run(
        docx_path=docx_path,
        images_dir=images_dir,
        heading_index=target["index"],
        threshold=0.55,
        apply=False,
        report_path=SCRATCH / "report_by_index.json",
    )
    check("heading_index run finds the same section as heading-text run",
          idx_report["section_heading"] == "Appendix B: STC and Erosion Maps"
          and len(idx_report["replacements"]) == 2 and len(idx_report["insertions"]) == 1)

    print()
    if failures:
        print(f"{len(failures)} check(s) FAILED:")
        for f in failures:
            print(f"  - {f}")
        sys.exit(1)
    else:
        print("All checks passed.")


if __name__ == "__main__":
    main()
